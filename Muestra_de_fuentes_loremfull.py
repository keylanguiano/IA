from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert
import fitz  # PyMuPDF
from PIL import Image, ImageOps
import os
import numpy as np

# Definir el texto de ejemplo
texto = (
    "Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. "
    "Árboles, niños, y jóvenes juegan en el jardín. "
    "¡Qué rápido corren los zorros! "
    "El pingüino Wenceslao hizo kilómetros bajo exhaustiva lluvia y frío. "
    "¿Cómo lograré sobrevivir sin mi teléfono móvil? "
    "1 2 3 4 5 6 7 8 9 0 "
    "A B C D E F G H I J K L M N O P Q R S T U V W X Y Z "
    "a b c d e f g h i j k l m n o p q r s t u v w x y z "
    "Á É Í Ó Ú á é í ó ú "
    r"! $ % & ( ) * + , - . / : ; < = > ? @ [ ] ^ _ ` { | } ~ "
)

texto_alternativo = (
    "Cada vez que subo a la azotea, miro lejos: montañas, ríos, la bruma infinita. "
    "Juguetes, risas y abrazos alegran el parque. "
    "¡Cuánto corre aquel jaguar bajo el sol! "
    "El búho majestuoso se esconde en ramas húmedas y frías. "
    "¿Quién encontrará la llave del misterio que guarda el viento? "
    "9 8 7 6 5 4 3 2 1 0 "
    "M Z X C V B N M A S D F G H J K L Ñ Q W E R T Y U I O P "
    "z x c v b n m a s d f g h j k l ñ q w e r t y u i o p "
    "Á É Í Ó Ú ü ä ë ï ö "
    r"~ @ # $ % ^ & * ( ) _ + : ; , - . / ? < > [ ] { } | ` ' = \" \\ "
)

# Lista de fuentes personalizadas
# 2
fonts = [
    "Abadi MT", "Agency FB", "Aharoni Bold", "Aldhabi", "Algerian", "Almanac MT", "American Uncial", "Andale Mono",
    "Andalus", "Andy", "Aparajita", "Arabic Transparent", "Arial",
    "Augsburger Initials", "Avenir Next LT Pro", "Bahnschrift", "Baskerville Old Face", "Batang & BatangChe", "Bauhaus 93",
    "Beesknees ITC", "Bell MT", "Bembo", "Berlin Sans FB", "Bernard MT Condensed", "Bickley Script", "Biome",
    "BIZ UDGothic", "BIZ UDMincho Medium", "Blackadder ITC", "Bodoni MT", "Bon Apetit MT",
    "Booakman Old Style", "Bookshelf Symbol", "Book Antiqua", "Bradley Hand ITC", "Braggadocio", "BriemScript",
    "Britannic Bold", "Broadway", "Brush Script MT", "Calibri", "Californian FB",
    "Calisto MT", "Cambria", "Cambria Math", "Candara", "Cariadings", "Centaur", "Century",
    "Century Gothic", "Century Schoolbook", "Chiller", "Colonna MT", "Comic Sans MS", "Consolas", "Constantia",
    "Contemporary Brush", "Cooper Black", "Copperplate Gothic", "Corbel", "CordiaUPC", "Courier New",
    "Curlz MT", "Dante", "DaunPenh", "David", "Daytona", "Desdemona", "Directions MT",
    "DokChampa", "Dotum & DotumChe", "Ebrima", "Eckmann", "Edda", "Elephant", 
    "Estrangelo Edessa", "Euphemia", "Eurostile", "FangSong", "Felix Titling",
    "Fine Hand", "Fixed Miriam Transparent", "Flexure", "Footlight MT", "Forte", "Franklin Gothic", "Franklin Gothic Medium",
    "FrankRuehl", "French Script MT", "Futura", "Gabriola", "Gadugi", "Garamond",
    "Garamond MT", "Gautami", "Georgia", "Georgia Ref", "Gigi", "Gill Sans MT", "Gill Sans MT Condensed", "Gisha",
    "Gloucester", "Goudy Old Style", "Gradl", "Grotesque", "Gulim & GulimChe", "Gungsuh & GungsuhChe",
    "Hadassah Friedlaender", "Harlow Solid Italic", "Harrington", "HGGothicE", "HGMinchoE",
    "HGSoeiKakugothicUB", "High Tower Text", "Holidays MT", "HoloLens MDL2 Assets", "Impact", "Imprint MT Shadow",
    "Informal Roman", "Iskoola Pota", "Jokerman", "Juice ITC", "KaiTi",
    "Kalinga", "Kartika", "Keystrokes MT", "Khmer UI", "Kigelia", "Kino MT", "Kokila", "Kristen ITC",
    "Lao UI", "Latha", "LCD", "Leelawadee", "Levenim MT", "Lucida Blackletter", "Lucida Bright",
    "Lucida Bright Math", "Lucida Calligraphy", "Lucida Console", "Lucida Fax", "Lucida Handwriting", "Lucida Sans",
    "Lucida Sans Typewriter", "Lucida Sans Unicode", "Magneto", "Maiandra GD", "Malgun Gothic", "Mangal", "Map Symbols", 
    "Matisse ITC", "Matura MT Script Capitals", "McZee", "Mead Bold", "Mercurius Script MT Bold",
    "Microsoft GothicNeo", "Microsoft JhengHei", "Microsoft JhengHei UI", "Microsoft New Tai Lue",
    "Microsoft PhagsPa", "Microsoft Sans Serif", "Microsoft Tai Le", "Microsoft YaHei",
    "Microsoft YaHei UI", "Microsoft Yi Baiti", "MingLiU", "MingLiU-ExtB", "MingLiU_HKSCS", "MingLiU_HKSCS-ExtB",
    "Minion Web", "Miriam", "Miriam Fixed", "Mistral", "Modern Love", "Modern No. 20", "Mongolian Baiti", "Monotype.com",
    "Monotype Corsiva", "MoolBoran", "MS Gothic", "MS LineDraw", "MS Mincho", "MS PGothic",
    "MS PMincho", "MS Reference", "MS UI Gothic", "MV Boli", "Myanmar Text", "Narkisim", "News Gothic MT",
    "New Caledonia", "Niagara", "Nirmala UI", "Nyala", "OCR-B-Digits", "OCRB", "OCR A Extended", "Old English Text MT",
    "Palatino Linotype", "Papyrus", "Parade", "Parties MT", "Peignot Medium",
    "Pepita MT", "Perpetua", "Perpetua Titling MT", "Placard Condensed", "Plantagenet Cherokee", "PMingLiU",
    "PMingLiU-ExtB", "Poor Richard", "Posterama", "Pristina", "Quire Sans", "Raavi", "Rage Italic", "Ransom", "Ravie",
    "RefSpecialty", "Rockwell", "Rockwell Nova", "Rod", "Runic MT Condensed", "Sabon Next LT", "Sagona", "Sakkal Majalla",
    "Script MT Bold", "Segoe Chess", "Segoe Print", "Segoe Script", "Segoe UI", "Segoe UI Symbol", "Selawik", "Shonar Bangla",
    "Showcard Gothic", "Shruti", "Signs MT", "SimHei", "Simplified Arabic Fixed", "SimSun", "SimSun-ExtB", "Sitka",
    "NSimSun", "Snap ITC", "Sports MT", "STCaiyun", "Stencil", "STFangsong", "STHupo", "STKaiti", "Stop", "STXihei",
    "STXingkai", "STXinwei", "STZhongsong", "Sylfaen", "Symbol", "Tahoma", "Tempo Grunge", "Tempus Sans ITC",
    "Temp Installer Font", "Times New Roman", "Times New Roman Special", "Tisa Offc Serif Pro",
    "Traditional Arabic", "Transport MT", "Trebuchet MS", "Tunga", "Tw Cen MT", "Univers", "Urdu Typesetting", "Utsaah",
    "Vacation MT", "Vani", "Verdana", "Verdana Ref", "Viner Hand ITC", "Vivaldi", "Vixar ASCI", "Vladimir Script",
    "Vrinda", "Walbaum", "Westminster", "Wide Latin"
]

fonts_extra = [
    "Engravers MT", "Enviro", "Angsana New", "Arial Black", "Arial Narrow", 
    "Arial Narrow Special", "Arial Rounded MT", "Arial Special", "Castellar",
    "Cavolini"
]

# 1
fonts_redimesion_to_small = [
    "Goudy Stout", "Javanese Text", "Meiryo", "Arial Unicode MS"
]

# 3
fonts_redimesion_to_medium = [
    "DFKai-SB", "Eras ITC"
]

# 4
fonts_redimesion_to_large = [
    "JasmineUPC", "Microsoft Uighur", "BrowalliaUPC", "AngsanaUPC"
    "Parchment", "Edwardian Script ITC", "Palace Script MT", "Bodoni MT Condensed", "IrisUPC", "The Hand", 
    "Haettenschweiler", "Onyx", "Vijaya", "Playbill", "Arabic Typesetting", "Cordia New", "The Serif Hand", "CordiaUPC",
    "LilyUPC", "Microsoft Himalaya", "Kunstler Script", "Freestyle Script", "FreesiaUPC", "Browallia New"
]

# 5
fonts_redimesion_to_xlarge = [
    "DilleniaUPC", "KodchiangUPC", "EucrosiaUPC"
]

# Lista de combinaciones de estilos
styles = [
    {"bold": False, "italic": False, "underline": False},
    {"bold": True, "italic": False, "underline": False},
    {"bold": False, "italic": True, "underline": False},
    {"bold": False, "italic": False, "underline": True},
    {"bold": True, "italic": True, "underline": False},
    {"bold": True, "italic": False, "underline": True},
    {"bold": False, "italic": True, "underline": True},
    {"bold": True, "italic": True, "underline": True}
]

# Crear un documento Word que contenga todas las combinaciones
doc = Document()

for font_name in fonts_redimesion_to_medium:
    for style_index, style in enumerate(styles):
        try:
            # Crear un párrafo y agregar el texto con el estilo indicado
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(texto * 3)

            # Configurar fuente y aplicar estilos
            run.font.name = font_name
            run.font.size = Pt(15)  # Ajuste del tamaño de la fuente para evitar desbordes
            run.bold = style["bold"]
            run.italic = style["italic"]
            run.underline = style["underline"]

            # Establecer alineación del párrafo
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            # Agregar un salto de página solo si no es la última combinación
            if not (font_name == fonts_redimesion_to_medium[-1] and style_index == len(styles) - 1):
                doc.add_page_break()
        
        except Exception as e:
            # Ignorar fuentes que no estén disponibles
            print(f"No se pudo usar la fuente {font_name}. Error: {e}")
            continue

# Guardar el documento de Word completo
output_word_path = "./SAMPLES/PROGRAM SAMPLES/GENERATE SAMPLES/Muestra_de_Fuentes.docx"
doc.save(output_word_path)
print(f"Archivo Word generado: {output_word_path}")

# Paso 2: Convertir el archivo de Word a PDF usando docx2pdf
output_pdf_path = "./SAMPLES/PROGRAM SAMPLES/GENERATE SAMPLES/Muestra_de_Fuentes.pdf"
convert(output_word_path, output_pdf_path)
print(f"Archivo PDF generado: {output_pdf_path}")

# Crear directorio para las imágenes recortadas
output_folder = "./SAMPLES/IMAGESs/TRAINING"
os.makedirs(output_folder, exist_ok=True)

# Paso 3: Convertir cada página del PDF a una imagen completa sin bordes
pdf_document = fitz.open(output_pdf_path)

# Recorrer cada página que no esté vacía y extraer toda el área de texto desde el primer renglón
for page_num in range(pdf_document.page_count):
    page = pdf_document[page_num]

    # Verificar si la página está vacía (ni texto ni elementos gráficos)
    pix = page.get_pixmap(dpi=72)
    img_np = np.array(Image.frombytes("RGB", [pix.width, pix.height], pix.samples))
    
    # Si la página está completamente en blanco (todos los píxeles son blancos), se omite
    if np.all(img_np == 255):
        continue

    # Obtener el pixmap de la página completa con mayor resolución
    pix = page.get_pixmap(dpi=1100)

    # Crear la imagen desde los bytes del pixmap
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

    # Crear fondo negro y agregar el texto en color blanco
    img_black = Image.new("RGB", img.size, (0, 0, 0))  # Crear un fondo negro
    img_white_text = ImageOps.invert(img)  # Invertir la imagen para tener texto blanco sobre fondo negro
    img_combined = Image.composite(img_white_text, img_black, mask=ImageOps.invert(img.convert("L")))

    # Recortar la imagen para eliminar los bordes en blanco
    bbox = img_combined.getbbox()
    if bbox:
        img_no_borders = img_combined.crop(bbox)
    else:
        img_no_borders = img_combined
        
    img_no_borders = ImageOps.expand (img_no_borders, border = 10, fill = 'black')

    # Redimensionar la imagen a un ancho de 1024 píxeles manteniendo la proporción
    target_width = 1024
    scale_factor = target_width / img_no_borders.width
    target_height = int(img_no_borders.height * scale_factor)

    img_resized = img_no_borders.resize((target_width, target_height), Image.LANCZOS)

    # Recortar la imagen a una altura de 512 píxeles
    crop_height = 512
    if img_resized.height > crop_height:
        img_cropped = img_resized.crop((0, 0, target_width, crop_height))
    else:
        img_cropped = img_resized  # Si la imagen es más pequeña que 512px de alto, no se recorta

    # Determinar la fuente y estilo basado en el número de página
    font_index = page_num // len(styles)
    style_index = page_num % len(styles)

    # Ajustar el índice de fuente para evitar desbordamientos
    if font_index >= len(fonts_redimesion_to_medium):
        break

    font_name = fonts_redimesion_to_medium[font_index].title ()
    font_name_safe = fonts_redimesion_to_medium[font_index].replace(" ", "_").title()

    # Generar el nombre del archivo según los estilos aplicados
    style = styles[style_index]
    style_name_parts = []
    if style["bold"]:
        style_name_parts.append("Negrita")
    if style["italic"]:
        style_name_parts.append("Cursiva")
    if style["underline"]:
        style_name_parts.append("Subrayado")
    style_name_safe = "_".join(style_name_parts) if style_name_parts else "Regular"

    # Generar la carpeta para la fuente y definir el nombre de la imagen
    font_folder = os.path.join(output_folder, font_name)
    os.makedirs(font_folder, exist_ok=True)

    img_name = os.path.join(font_folder, f"Loremfull_2_{font_name_safe}_{style_name_safe}.png")

    # Guardar la imagen con el nombre correspondiente
    img_cropped.save(img_name, format='PNG', quality=100, optimize=True, progressive=True)
    print(f"Imagen guardada: {img_name}")

# Cerrar el documento PDF
pdf_document.close()
print("Recortes completados.")
